# Source Generated with Decompyle++
# File: PPQ builder with GUI.pyc (Python 3.10)

import pandas as pd
import docx
from docx.shared import Cm
import easygui
import os
data = []
search = [
    [
        'Year'],
    [
        'Series'],
    [
        'TZ'],
    [
        'Topic'],
    [
        'Spec point']]
years = [
    '2017',
    '2018',
    '2019',
    '2020']
series = [
    'May',
    'November']
timezone = [
    '1',
    '2',
    '3']
topics = [
    '1 - States of matter',
    '2 - Atoms, elements and compounds',
    '3 - Stoichiometry',
    '4 - Electrochemistry',
    '5 - Chemical energetics',
    '6 - Chemical reactions',
    '7 - Acids, bases and salts',
    '8 - The Periodic Table',
    '9 - Metals',
    '10 - Chemistry of the environment',
    '11 - Organic chemistry',
    '12 - Experimental techniques and chemical analysis']
points = [
    '1.1 - Solids, liquids and gases',
    '1.2 - Diffusion',
    '2.1 - Elements, compounds and mixtures',
    '2.2 - Atomic structure and the Periodic table',
    '2.3 - Isotopes',
    '2.4 - Ions and Ionic bonds',
    '2.5 - Simple molecules and covalent bonds',
    '2.6 - Giant covalent structures',
    '2.7 - Metallic bonding',
    '3.1 - Formulae',
    '3.2 - Relative masses of atoms and molecules',
    '3.3 - The mole and the Avogadro constant',
    '4.1 - Electrolysis',
    '4.2 - Hydrogen\xe2\x80\x93oxygen fuel cells',
    '5.1 - Exothermic and endothermic reactions',
    '6.1 - Physical and chemical changes',
    '6.2 - 6.1 Physical and chemical changes',
    '6.3 - Reversible reactions and equilibrium',
    '6.4 - Redox',
    '7.1 - The characteristic properties of acids and bases',
    '7.2 - Oxides',
    '7.3 - Preparation of salts',
    '8.1 - Arrangement of elements',
    '8.2 - Arrangement of elements',
    '8.3 - Group VII properties',
    '8.4 - Transition elements',
    '8.5 - Noble gases',
    '9.1 - Properties of metals',
    '9.2 - Uses of metals',
    '9.3 - Alloys and their properties',
    '9.4 - Reactivity series',
    '9.5 - Corrosion of metals',
    '9.6 - Extraction of metals',
    '10.1 - Water',
    '10.2 - Fertilisers',
    '10.3 - Air quality and climate',
    '11.1 - Formulae, functional groups and terminology',
    '11.2 - Naming organic compounds',
    '11.3 - Fuels',
    '11.4 - Alkanes',
    '11.5 - Alkenes',
    '11.6 - Alcohols',
    '11.7 - Carboxylic acids',
    '11.8 - Polymers',
    '12.1 - Experimental design',
    '12.2 - Acid\xe2\x80\x93base titrations',
    '12.3 - Chromatography',
    '12.4 - Separation and purification',
    '12.5 - Identification of ions and gases']

def csv_converter():
    data1 = pd.read_excel('PPQ directory.xlsx')
    df = pd.DataFrame(data1)
    data_list = df.values.tolist()
    return data_list


def filter_images():
    for a in range(1, len(data)):
        topic_list = []
        sp_list = []
        spec_points = data[a][6]
        spl_spec_points = spec_points.split(';')
        for e in range(len(spl_spec_points)):
            spec_point = spl_spec_points[e]
            sp_list.append(spec_point)
            spec_point_spl = spec_point.split('.')
            topic = spec_point_spl[0]
            topic_list.append(topic)

        year = 0
        series = 0
        tz = 0
        if str(data[a][0]) in search[0] or search[0][1] == '':
            year += 1
        if data[a][1] in search[1] or search[1][1] == '':
            series += 1
        if str(data[a][2]) in search[2] or search[2][1] == '':
            tz += 1
        unit = 0
        for b in range(len(topic_list)):
            if topic_list[b] in search[3]:
                unit += 1
            if search[3][1] == '':
                unit += 1
        sp = 0
        for c in range(len(sp_list)):
            spec = sp_list[c]
            spec_split = spec.split('.')
            spec_theme = spec_split[0] + '.' + spec_split[1]

            if spec_theme in search[4]:
                sp += 1
            if search[4][1] == '':
                sp += 1

        if year > 0 and series > 0 and tz > 0 and unit > 0 and sp > 0:
            folder = os.getcwd()
            file = os.path.join(folder, 'File Bank', data[a][4] + '.png')
            ms = os.path.join(folder, 'MS Bank', data[a][5] + '.png')
            images.append(file)
            markscheme.append(ms)


def write_paper(docname):
    document1 = docx.Document()
    document1.add_heading(docname, 0)
    sections = document1.sections
    section = sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1.5)

    for b in range(len(images)):
        question = str(b + 1)
        document1.add_paragraph(question)
        document1.add_picture(images[b], width=Cm(4.5*2.54))

    try:
        document1.save(name + '.docx')
    except TypeError:
        pass
    


def write_ms(name):
    document2 = docx.Document()
    document2.add_heading(name, 0)
    sections = document2.sections
    section = sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1.5)

    for b in range(len(markscheme)):
        question = str(b + 1)
        document2.add_paragraph(question)
        document2.add_picture(markscheme[b], width=Cm(4.5*2.54))

    try:    
        document2.save(name + ' MS.docx')
    except TypeError:
        pass

data = csv_converter()

choices = [
    'By Year',
    'By Topic',
    'By Syllabus Point',
    'All']

reply = easygui.buttonbox('iGCSE PPQs', 'Choose how you want to filter the PPQs', choices, 'chem.png',)

if reply == 'By Year':
    try:
        byyear = easygui.multchoicebox('By Year', 'By Year', years)
        for x in range(len(byyear)):
            search[0].append(byyear[x])
    except TypeError:
        pass

    search[1].append('')
    search[2].append('')
    search[3].append('')
    search[4].append('')
    name = easygui.enterbox("Specify the name of the output file (No .docx needed)", 'Filename')
    
elif reply == 'By Topic':
    bytopic = easygui.multchoicebox('By Topic', "By Topic", topics)
    try:
        for x in range(len(bytopic)):
            numpointlist = bytopic[x].split(' -')
            numpoint = numpointlist[0]
            search[3].append(numpoint)
    except TypeError:
        pass

    search[0].append('')
    search[1].append('')
    search[2].append('')
    search[4].append('')
    name = easygui.enterbox("Specify the name of the output file (No .docx needed)", 'Filename')


elif reply == 'By Syllabus Point':
    bypoint = easygui.multchoicebox('By Syllabus Point', 'By Syllabus Point', points)
    
    try:
        for x in range(len(bypoint)):
            numpointlist = bypoint[x].split(' -')
            numpoint = numpointlist[0]
            search[4].append(numpoint)
    except TypeError:
        pass

    search[0].append('')
    search[1].append('')
    search[2].append('')
    search[3].append('')
    name = easygui.enterbox("Specify the name of the output file (No .docx needed)", 'Filename')

elif reply == 'All':
    search[0].append('')
    search[1].append('')
    search[2].append('')
    search[3].append('')
    search[4].append('')
    name = 'All'

images = []
markscheme = []

filter_images()

write_paper(name)
write_ms(name)